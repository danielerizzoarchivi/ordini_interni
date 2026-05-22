import os
import zipfile
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_TEMPLATE = os.path.normpath(os.path.join(
    os.path.dirname(__file__), '..', '..',
    'ordini_interni', 'Carta Intestata.dotx'
))


def _open_template():
    """Apre il .dotx come Document patchando il content type."""
    buf = BytesIO()
    with zipfile.ZipFile(_TEMPLATE, 'r') as zin:
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    data = data.replace(
                        b'wordprocessingml.template.main+xml',
                        b'wordprocessingml.document.main+xml'
                    )
                zout.writestr(item, data)
    buf.seek(0)
    return Document(buf)


def _clear_body(doc):
    """Svuota il body preservando sectPr (margini, header/footer)."""
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    for child in list(body):
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, size=10, align=None, color=None):
    cell.text = ''
    para = cell.paragraphs[0]
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run


def _fmt_eur(value):
    if value is None:
        return ''
    return f'€ {float(value):,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')


def genera_docx(ordine, cfg):
    doc = _open_template()
    _clear_body(doc)

    # --- Destinatario ---
    # (l'intestazione aziendale è nel page header del template)
    doc.add_paragraph('Spett.le').runs[0].font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run(ordine.fornitore_nome.upper())
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    # --- Tabella header ordine ---
    t_header = doc.add_table(rows=2, cols=6)
    t_header.style = 'Table Grid'

    row0 = t_header.rows[0]
    _cell_text(row0.cells[0], 'Ordine n.', bold=True)
    _cell_text(row0.cells[1], ordine.numero_completo, bold=True)
    row0.cells[1].merge(row0.cells[2])
    row0.cells[1].merge(row0.cells[3])
    _cell_text(row0.cells[4], 'Data:', bold=True)
    _cell_text(row0.cells[5], ordine.creato_il.strftime('%d/%m/%Y'))

    row1 = t_header.rows[1]
    _cell_text(row1.cells[0], 'Emesso da:', bold=True)
    nome_emittente = ordine.emesso_da.nome_completo if ordine.emesso_da else ''
    _cell_text(row1.cells[1], nome_emittente)
    row1.cells[1].merge(row1.cells[2])
    row1.cells[1].merge(row1.cells[3])
    _cell_text(row1.cells[4], 'Rif.:', bold=True)
    _cell_text(row1.cells[5], ordine.riferimento or '')

    widths = [Cm(2.5), Cm(5), Cm(0.1), Cm(0.1), Cm(1.5), Cm(3)]
    for i, w in enumerate(widths):
        for row in t_header.rows:
            row.cells[i].width = w

    doc.add_paragraph()

    # --- Tabella articoli ---
    t_art = doc.add_table(rows=1, cols=4)
    t_art.style = 'Table Grid'

    hdr = t_art.rows[0]
    for i, (txt, al) in enumerate([('Descrizione', None), ('Q.tà', 'center'),
                                    ('P.un', 'right'), ('Prezzo Totale', 'right')]):
        _set_cell_bg(hdr.cells[i], '1F4E79')
        _cell_text(hdr.cells[i], txt, bold=True, color='FFFFFF', align=al)

    for art in ordine.articoli:
        row = t_art.add_row()
        _cell_text(row.cells[0], art.descrizione)
        _cell_text(row.cells[1], f'{float(art.quantita):g}'.replace('.', ','), align='center')
        _cell_text(row.cells[2], _fmt_eur(art.prezzo_unitario), align='right')
        _cell_text(row.cells[3], _fmt_eur(art.totale), align='right')

    sconto = ordine.sconto_totale
    if sconto:
        row = t_art.add_row()
        _cell_text(row.cells[0], 'SCONTO', bold=True)
        row.cells[0].merge(row.cells[1])
        row.cells[0].merge(row.cells[2])
        _cell_text(row.cells[3], f'- {_fmt_eur(sconto)}', bold=True, align='right')

    row = t_art.add_row()
    _set_cell_bg(row.cells[0], 'DDEBF7')
    _cell_text(row.cells[0], 'TOTALE', bold=True)
    row.cells[0].merge(row.cells[1])
    totale_testo = f'{_fmt_eur(ordine.totale_netto)} (IVA INCLUSA)'
    _cell_text(row.cells[2], totale_testo, bold=True, align='right')
    _set_cell_bg(row.cells[2], 'DDEBF7')
    _cell_text(row.cells[3], totale_testo, bold=True, align='right')
    _set_cell_bg(row.cells[3], 'DDEBF7')

    art_widths = [Cm(9), Cm(1.8), Cm(2.5), Cm(3)]
    for w_col, w in enumerate(art_widths):
        for row in t_art.rows:
            row.cells[w_col].width = w

    doc.add_paragraph()

    # --- Note ---
    note_p = doc.add_paragraph()
    run = note_p.add_run(
        'IMPORTANTE: RIPORTARE IL N° ORDINE SU DDT E FATTURE. '
        'NON SARANNO LIQUIDATE LE FATTURE SENZA IL N° ORDINE'
    )
    run.bold = True
    run.font.size = Pt(9)

    if ordine.pagamento:
        doc.add_paragraph(f'PAGAMENTO: {ordine.pagamento}').runs[0].font.size = Pt(10)
    if ordine.consegna:
        doc.add_paragraph(f'CONSEGNA: {ordine.consegna}').runs[0].font.size = Pt(10)
    if ordine.note:
        doc.add_paragraph(ordine.note).runs[0].font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def genera_pdf(ordine, app):
    from weasyprint import HTML
    from flask import render_template

    with app.app_context():
        html = render_template('orders/print.html', ordine=ordine, cfg=app.config)

    pdf = HTML(string=html, base_url=app.static_folder).write_pdf()
    return pdf
